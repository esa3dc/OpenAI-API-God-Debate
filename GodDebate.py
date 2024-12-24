import openai 
import time
from docx import Document
import re

# Set your OpenAI API key
# replace 'KEY' with API key
openai.api_key = 'KEY'

# Persona definitions
theist_persona = "You are Fyodor Dostoevsky. You believe in God and are having a conversation about God's existence."
atheist_persona = "You are Richard Dawkins. You don't believe in God and  are having a conversation about God's existence."

# Debate configuration
debate_topic = "Does God exist?"
rounds = 3  # Number of debate rounds

def get_response(prompt, persona, model="gpt-4-turbo", temp = 0.7):
    """Send a prompt to the LLM model with a specific persona and get the response."""
    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": persona},
                {"role": "user", "content": prompt}
            ],
            temperature = temp
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        print(f"Error: {e}")
        return "Sorry, there was an error generating the response."

def debate():
    doc = Document()
    doc.add_heading(f"Debate On: {debate_topic}\n", 0)
    prompt = f"The conversation topic is: {debate_topic}\nDostoevsky starts:"
    

    # Initial prompt to start the debate
    for i in range(rounds):
        # Theist argument
        print(f"Round {i+1} begins:")
        theist_response = get_response(prompt, theist_persona)
        cleaned_theist_response = re.sub(r'[^\x20-\x7E]+', '', theist_response)
        doc.add_paragraph(f"Round {i+1} - Dostoevsky: {cleaned_theist_response}")

        # Update the prompt with the theist's argument for the atheist to respond
        prompt = f"Dostoevsky said: {cleaned_theist_response}\nDawkins, what is your response?"
        
        # Short delay for realistic effect
        time.sleep(1)

        # Atheist argument
        atheist_response = get_response(prompt, atheist_persona)
        cleaned_atheist_response = re.sub(r'[^\x20-\x7E]+', '', atheist_response)
        doc.add_paragraph(f"Round {i+1} - Dawkins: {cleaned_atheist_response}")

        # Update the prompt with the atheist's argument for the theist to respond
        prompt = f"Dawkins said: {cleaned_atheist_response}\nDostoevsky, what is your response?"
        
        # Another short delay
        time.sleep(1)

    # Final remarks
    print("\nDebate concluded.")
    doc.add_paragraph("\nDebate concluded.")
    doc.save("Debate_Output.docx")

# Run the debate
debate()
